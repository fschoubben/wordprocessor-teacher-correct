import os
from docx import Document
from PIL import Image
import io
import math
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def check_sections_word(doc, student, key="section"):
    # CAUTION : these are "bonus points"
    if len(doc.sections) > 1:
        student.scores[key] = student.max_points[key]
    else:
        # CAREFUL: these are "bonus points"
        student.scores[key] = 0
        student.reasons[key] = "pas de section trouvée. "

def check_quote(doc, student, key="section", debug=False):
    max_score = student.max_points[key]
    score=0
    for para in doc.paragraphs:
        #print_debug(debug, str(para))
        if has_quote(para, debug):
            score = max_score
            break
    if score==0:
        student.reasons[key] = "pas de citation trouvée. "
    student.to_check_manually += "vérifier citation avec note de bas de page"
    student.scores[key] = score

def has_quote(para, debug=False):
    text = para.text
    #print(text)
    start_index = text.find("«")
    if start_index != -1:
        print_debug(debug, "« found")
        end_index = text.find("»", start_index + 1)
        if end_index != -1:
            print_debug(debug, "» found")
            quote_text = text[start_index + 1:end_index]  # Extract text between French quotes
            # TODO : check footnote after
            return True
    start_index = text.find("«")
    if start_index!=-1:
        print("étrange...")
    start_index = text.find('"')
    if start_index != -1:
        print_debug(debug, "\" 1 found")
        end_index = text.find('"', start_index + 1)
        if end_index != -1:
            print_debug(debug, "\" 2 found")
            quote_text = text[start_index + 1:end_index]  # Extract text between French quotes
            # TODO : check footnote after
            return True
    return False

def check_picture_proportions(doc):
    all_proportions_maintained = True

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run._element.xml.startswith('<w:drawing'):
                shape = run._element
                if shape.type == 3:  # Picture shape type
                    aspect_ratio = get_aspect_ratio(shape)
                    if not math.isclose(aspect_ratio, 1.0, rel_tol=1e-5):  # Assuming original aspect ratio is 1:1
                        all_proportions_maintained = False
                        break

    return all_proportions_maintained

def get_aspect_ratio(shape):
    blip = shape.xpath('.//a:blip', namespaces=shape.nsmap)
    if blip:
        embed = blip[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        rel = shape.part.rels.get(embed)
        if rel:
            image = rel.target_part.blob
            width, height = Image.open(io.BytesIO(image)).size
            return width / height

    return 1.0  # Default aspect ratio if unable to retrieve
def print_debug(debug, message):
    if debug:
        print(message)

if __name__ == "__main__":
    #filename = "2023-01-TIC1-Test-1.docx"
    #filename = "2023-01-TIC1-Test-2.docx"
    filename = "2024-01-S2-Test-8.docx"
    # TODO : why check_quote don't work with Test 8 ??
    # TODO : rewrite as a macro, there are too much strange things :-(
    # filename = "2024-01-S2-Test-8.docx"
    from student import Student
    stud = Student()
    # Créer une instance de l'application Word
    #word_app = win32com.client.Dispatch("Word.Application")

    # Ouvrir le document

    #document = word_app.Documents.Open(document_path)

    print(filename)
    path = os.getcwd()
    filename = path+'/'+filename
    print(filename)

    document = Document(filename)
    check_quote(document, stud, key="citation", debug=False)
    if stud.scores["citation"]==0:
        print("pas de citation")
    else:
        print("citation OK")
    proportions_ok = check_picture_proportions(document)
    if proportions_ok:
        print("proportions OK")
    else:
        print("problème de proportions")