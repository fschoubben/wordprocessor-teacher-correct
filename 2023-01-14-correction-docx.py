import sys
import os
import psutil
from tkinter import messagebox

import PyPDF2
from word_macros import *
from word_functions import *

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import win32com.client

import openpyxl
# from openpyxl import load_workbook
from openpyxl.comments import Comment
from openpyxl.utils.cell import get_column_letter
from openpyxl.styles import PatternFill
from openpyxl.styles import Font

from student import Student

debug = True

header_to_check="S2"
middle_footer_to_check="S2-B1 - Numérique"
excel_file_for_results = "./2024-01-auto-correct-results.xlsx"
default_start_of_filename = "2024-01-S2-"

# TODO : Check if check_table works (seems to give 2 points when no file exists !?
# TODO  everything in english
# TODO  Later : internationalisation

# TODO  vérifier que c'est bien un fichier Word ou ... avant de lancer le reste : for el in *.docx? filetype ?
# TODO  utilisation student systématique

# TODO : vérifier si un fichier est déjà ouvert avant de l'ouvrir, en Word. Prévenir de le fermer si oui.
# TODO : would it be better to force close if a file is already open ? For xlsx, maybe.
# TODO : ajouter mesure temps par fonction, parce que c'est trop lent (commencer par returns ! )

# TODO : Formats
# todo : orthographe
# todo : citation
# todo : note bas de page

def execute_ensuring_file_not_open(file, command):
    command_executed = False
    while not command_executed:
        if os.path.exists(file):
            try:
                os.rename(file, file)
                print_debug(debug, 'Access on file "' + file + '" is available!')
                command(file)
                command_executed = True
            except OSError as e:
                print('Access-error on file "' + file + '"! \n' + str(e))
                messagebox.showinfo(title="Script de correction automatique",
                                    message="Fermer le document Excel pour la nouvelle correction")
        else:
            command(file)
    # workbook.save(excel_file_for_results)


def fill_first_lines_excel(worksheet, student):
    row = 1
    worksheet.cell(row=row, column=1).value = "Nom"
    worksheet.cell(row=row, column=1).font = Font(bold=True)
    worksheet.cell(row=row, column=2).value = "Prénom"
    worksheet.cell(row=row, column=2).font = Font(bold=True)
    worksheet.cell(row=row, column=3).value = "Total"
    worksheet.cell(row=row, column=3).font = Font(bold=True)
    col = 4
    for key in student.scores.keys():
        worksheet.cell(row=row, column=col).value = key
        worksheet.cell(row=row, column=col).font = Font(bold=True)
        col += 1
    worksheet.cell(row=row, column=col).value = "à vérifier manuellement"
    worksheet.cell(row=row, column=col).font = Font(bold=True)
    col = 4
    row += 1
    for key, value in student.max_points.items():
        worksheet.cell(row=row, column=col).value = key
        worksheet.cell(row=row, column=col).font = Font(bold=True)
        worksheet.cell(row=row, column=col).font = Font(italic=True)
        col += 1
    row += 1
    col = 4
    for key, value in student.max_points.items():
        worksheet.cell(row=row, column=col).value = value
        worksheet.cell(row=row, column=col).font = Font(bold=True)
        col += 1
    worksheet.cell(row=row, column=3).value = "=sum(" + \
                                              get_column_letter(4) + str(row) + \
                                              ":" + get_column_letter(4 + len(student.scores.items())) + \
                                              str(row) + ")"

    worksheet.cell(row=row, column=col).font = Font(bold=True)

    row += 1

    worksheet.freeze_panes = 'D4'
    # col += 2
    # for key in reasons_set:
    #     worksheet.cell(row=1, column=col).value = key
    #     col += 1
    return row


def fill_result_line_in_excel(worksheet, row, student):
    # Pour chaque élément du set, l'ajouter dans une nouvelle cellule
    worksheet.cell(row=row, column=1).value = student.name.capitalize()
    worksheet.cell(row=row, column=2).value = student.firstname.capitalize()
    worksheet.cell(row=row, column=3).value \
        = "=sum(" + get_column_letter(4) + str(row) + ":" + get_column_letter(4 + len(student.scores.items())) + str(
        row) + ")"
    col = 4
    # print("à vérifier : ", student.to_check)
    # print("sytles : ", student.scores["styles"], "et liens : ", student.scores["lien"])
    # TODO add conditional formatting : https://openpyxl.readthedocs.io/en/latest/formatting.html
    #        --> < max_score/2 --> red font color
    #        --> = max_score   --> green font color
    # Todo add formulas : https://openpyxl.readthedocs.io/en/latest/usage.html?highlight=formula#using-formulae
    for key, value in student.scores.items():
        worksheet.cell(row=row, column=col).value = value
        cell = get_column_letter(col) + str(row)
        blank_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
        worksheet[cell].fill = blank_fill
#was        blank_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
        # Create the comment with the "Why that score" if any
        why = student.reasons[key]
        if why != "":
            comment = Comment(why, "François Schoubben")
            worksheet[cell].comment = comment
        if key in student.to_check:
            # print("mettre", key, " en jaune")
            yellow_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
            worksheet[cell].fill = yellow_fill
        col += 1

    worksheet.cell(row=row, column=col).value = student.to_check_manually


def fill_last_line_in_excel(worksheet, row, student, number_of_non_student_lines):
    row += 1
    worksheet.cell(row=row, column=1).value = "Moyenne étudiant"
    for col in range(3, len(student.max_points) + 3):
        worksheet.cell(row=row, column=col).value = \
            "=average(" + get_column_letter(col) + str(number_of_non_student_lines + 1) \
            + ":" + get_column_letter(col) + str(row - 1) + ")"

    row += 1
    worksheet.cell(row=row, column=1).value = "Min étudiant"
    for col in range(3, len(student.max_points) + 3):
        worksheet.cell(row=row, column=col).value = \
            "=min(" + get_column_letter(col) + str(number_of_non_student_lines + 1) \
            + ":" + get_column_letter(col) + str(row - 2) + ")"
    row += 1
    worksheet.cell(row=row, column=1).value = "MAX étudiant"
    for col in range(3, len(student.max_points) + 3):
        worksheet.cell(row=row, column=col).value = \
            "=max(" + get_column_letter(col) + str(number_of_non_student_lines + 1) \
            + ":" + get_column_letter(col) + str(row - 3) + ")"

# # gestion des fichiers
# nom fichier :  2023-01-TIC1—Nom-


def verifier_nom_fichiers(mfile, template, student):
    # TODO NextVersion (manually done 2023): vérifier espaces ! (-1)
    raisons = ""
    points = 0
    if mfile.startswith(template):
        points += student.max_points["nomFichiers"]
    #mfile = mfile.replace("—", "")
    #mfile = mfile.replace(" ", "")
    # if "schoubben" in mfile:  # TODO 2023 : manually done, try to emprove to automate...
    #     raisons="mauvais nom de fichier; "
    nomcomplet = mfile.split("-")
    nom = nomcomplet[3]
    prenom = nomcomplet[4]
    # print("in nom prenom = ", nom, prenom, nomcomplet)
    student.name = nom
    student.firstname = prenom
    student.scores["nomFichiers"] = points
    student.reasons["nomFichiers"] = raisons
    return student.max_points["nomFichiers"]


# 2 formats : word/opendoc + pdf /2 TODO : améliorer pour vérifier types
def verifier_deux_formats_fichiers(filename, liste_fichiers, max_points, scores_set, reasons_set, key="format"):
    # f : filename (string)
    nb_fichiers = 0
    for el in liste_fichiers:
        if filename[0:-4] in el:
            nb_fichiers += 1
    # print("nb fichiers", nbFichiers)
    if nb_fichiers == 2:
        scores_set[key] = max_points
        reasons_set[key] = ""
    else:
        scores_set[key] = 0
        reasons_set[key] = "il n'y a pas les 2 formats de fichier"


# moins de 3Mo (moins de 1Mo) /2
def verifier_moins_de_3_mo(filename, max_size, max_points):
    # f : filename (string)
    try:
        file_info = os.stat(filename)
        if file_info.st_size < max_size * 1000000:
            return max_points, ""
        else:
            return 0, "fichier trop gros"
    except Exception as e:
        sys.stderr.write("erreur de nom de fichier" + str(e))


# Traitement de texte
# min 3, max 10 pages /2


def verifier_nombre_pages_pdf(pdf, min_pages, max_pages, student):
    """ input : f = fichier pdf existant,
    # min = nombre de page minimum (>=0),
    # max = nombre de pages max (>=0),
    # pts = nombre de points à attribuer
    # output : pts si le nombre de page de f est compris dans l'intervale [min, max], 0 sinon
    # Obtenez les informations du document"""
    # info = pdf.metadata
    key = "pages"
    doc_pages_count = len(pdf.pages)
    if min_pages <= doc_pages_count <= max_pages:
        student.scores[key] = student.max_points[key]
        student.reasons[key] = ""
    else:
        if len(pdf.pages) > max_pages:
            student.scores[key] = 0
            student.reasons[key] = "Plus de " + str(max_pages) + " pages"
        else:
            student.scores[key] = 0
            student.reasons[key] = "Moins de " + str(min_pages) + " pages"
    return doc_pages_count


# orthographe OK (/0) ?? grammalect ? TODO ?
# utilisation de styles : Titre 1, Titre 2, normal == justifié
def lister_styles_word(doc):
    listeTitres = []
    #     for parag in doc.paragraphs:
    #         if ("Heading" in parag.style.name) or ("Titre" in parag.style.name):
    #             #print("oui")
    #             listeTitres.append(parag)
    styles = set()
    for paragraph in doc.paragraphs:
        # Obtenez le nom du style du paragraphe
        style_name = paragraph.style.name
        # Si le nom du style correspond à un style de titre, imprimez le contenu du paragraphe
        styles.add(style_name)
        if style_name in ('Titre 1', 'Titre 2', 'Titre 3', 'Heading 1', 'Heading 2', 'Heading 3'):
            listeTitres.append(paragraph)
    #print(styles)
    return listeTitres


def verifier_styles_word(doc, student, key = "styles"):
    pts = 0
    raison = ""

    # utilisation styles pour titres
    lstyles = lister_styles_word(doc)
    # print("nombre de titres utilisés :", len(lstyles))
    if len(lstyles) > 0:
        pts = student.max_points[key] / 2
    else:
        raison += "Pas de style de titre utilisé. "
    # for el in lstyles:
    #    print(el.style_name, el.text)

    # justifié par le style normal
    normal_style = doc.styles['Normal']
    # print(normal_style.paragraph_format.alignment)
    if normal_style.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
        pts += student.max_points[key] / 2
    else:
        raison += "Le style normal n'est pas justifié. "
        student.to_check_manually += "est-ce que cest bien le style normal qui est utilisé ? si non, est-ce justifié ? "
        student.to_check.add(key)

    student.scores[key] = pts
    student.reasons[key] = raison


def check_page_number_Word(doc, total_pages):
    print("Nombre total de pages (Word) : ", total_pages)
    # Récupérer les pieds de page
    footer_text = doc.Sections.Item(1).Footers.Item(1).Range.Text
    if str(total_pages) in footer_text:
        print(f"Le nombre total de pages est présent dans le pied de page : \" ", footer_text, " \"")
    else:
        print("pas de nombre total de pages dans ", footer_text)

    # footers = doc.Sections.Item(1).Footers
    # for footer in footers:
    #     for shape in footer.Shapes:
    #         if shape.Type == 8:
    #             if shape.TextFrame.TextRange.Text.find("{ NUMPAGES }") != -1:
    #                 print("This document contains a field with total number of pages in footer.")
    #                 break
    # else:
    #     print("This document doesn't contain a field with total number of pages in footer.")


# en-tête / pied de page /4
# en haut à droite votre section (NP ou NPS ou NS),
# en bas à gauche votre nom,
# en bas à droite le numéro de page et le nombre de pages,
# en bas et au centre « Examen TICE — B1 » ;


# espaces / sauts de page :
#   max 2 enter, max 2 espaces d'affilée, bonus pour style de page/saut de section
def verifier_nombre_enter_et_espaces_word(doc, student, max_return=3, max_spaces=2, key ="espaces"):  #
    # doc == pydocx
    points = student.max_points[key]
    previous_empty = False
    espacesTrouves = False
    raison_enter = ""
    raison_spaces = ""

    paragraphs = doc.paragraphs
    consecutiveemptyparagraphscount = 0

    for p in paragraphs:
        content=p.text
        content.replace(" ", "")
        content.replace("\t", "")
        if len(content) <= 1:
            consecutiveemptyparagraphscount += 1
            if consecutiveemptyparagraphscount > max_return:
                #more than max_enter return founds
                break
        else:
            consecutiveemptyparagraphscount = 0

    if consecutiveemptyparagraphscount > max_return:
        points -= student.max_points[key] / 2
        raison_enter = "Il y a plusieurs retours à la ligne consécutifs dans le document. "

    for paragraph in doc.paragraphs:
        too_many_spaces=" "*max_spaces
        if too_many_spaces in paragraph.text:
            ##print("plusieurs espaces")
            if not espacesTrouves:
                points -= student.max_points[key] / 2
                espacesTrouves = True
                raison_spaces = "Il y a plusieurs espaces consécutives dans le document. "
    if points <= 0:
        student.scores[key] = 0
        student.reasons[key] = raison_spaces + raison_enter
        # TODO : créer une classe "fichier word" (objet + méthodes ?)
    else:
        student.scores[key] = points
        if points < student.max_points[key]:
            student.reasons[key] = raison_spaces + raison_enter


def check_page_returns_word(document, student, key):
    # document : win32com
    range = document.Content
    range.Find.ClearFormatting()
    range.Find.Replacement.ClearFormatting()
    range.Find.Text = ""
    range.Find.Format = False
    range.Find.Forward = True
    range.Find.Wrap = 1
    range.Find.Execute(FindText=chr(12))

    if range.Find.Found:
        if student.scores[key] < 3:
            student.scores[key] += 1
            student.reasons[key] += "Il y a des sauts de page"

# listes et sous-listes : /2
def verifier_listes_word(document, student, key):
    listes = False
    sous_listes = False

    # Pour chaque paragraphe du document
    for paragraph in document.Paragraphs:
        # Si le paragraphe est dans une liste
        if paragraph.Range.ListFormat.ListType != 0:
            listes = True
            # Si le paragraphe est dans une sous-liste
            if paragraph.Range.ListFormat.ListLevelNumber > 1:
                sous_listes = True
                break

    if sous_listes:
        student.scores[key] = student.max_points[key]
        student.reasons[key] = ""
    elif listes:
        student.scores[key] = student.max_points[key] / 2
        student.reasons[key] = "Pas de sous-liste"
    else:
        student.scores[key] = 0
        student.reasons[key] = "Pas de liste"


# citation avec guillemets : /2 TODO ??
# note de bas de page avec sources : / 2 TODO
def check_has_footnotes_word(doc, student, key = "noteBasPage"):
    footnotes = doc.Footnotes
    if footnotes.Count > 0:
        student.scores[key] = student.max_points[key]
        student.reasons[key] = ""
        # print("This document contains footnote(s)")
    else:
        endnotes = doc.Endnotes
        if endnotes.Count > 0:
            student.scores[key] = student.max_points[key]
            student.reasons[key] = "Sous forme de note de fin de document. "
        else:
            student.scores[key] = 0
            student.reasons[key] = "Pas de note de bas de page. "
        # print("This document doesn't contain any footnote")


def verifier_TDM_word(document, student):
    key="TDM"
    if document.TablesOfContents.Count > 0:
        student.scores[key] = student.max_points[key]
        student.reasons[key] = ""
    else:
        student.scores[key] = 0
        student.reasons[key] = "pas de table des matières"


# image : /4 TODO
#	redimenssionnée avec proportions ok
#	légende
#	texte à gauche de l'image
#	espacement de min 0,5 avec le texte
# def verifier_images_redimensionnees_correctement_PDF(file, pdf_reader, maxPoints):
#     pages = pdf_reader.getNumPages()
#
#     for page_number in range(pages):
#         page = pdf_reader.getPage(page_number)
#         for obj in page['/Resources']['/XObject'].values():
#             if obj['/Subtype'] == '/Image':
#                 if '/FlateDecode' in obj['/Filter']:
#                     print("This image is PNG")
#                 elif '/DCTDecode' in obj['/Filter']:
#                     print("This image is JPG")


def verifier_images_redimensionnees_correctement_word(document, student, key="images"):
    # Pour chaque image du document
    raisons = ""
    max_points = student.max_points[key]
    to_check_manually = "vérifier si image a gardé proportions + texte à gauche + légende"
    # TODO : vérifier le reste
    resize_OK = True
    points = 0
    has_image = False
    # if len(document.Shapes)>0:
    #     points+=maxPoints/4
    image_count = document.Shapes.Count
    # print("nombre d'images \"Shape\" : ", image_count)
    if image_count > 0:
        has_image = True
        for shape in document.Shapes:
            if shape.Type == 12:  # image
                if shape.HasCaption:
                    print(f"shape : {shape.Name} has a caption")
                else:
                    print(f"shape : {shape.Name} doesn't have a caption")
    image_count = document.InlineShapes.Count
    if image_count > 0:
        has_image = True

    if has_image:
        points += max_points / 4
    else:
        raisons += "Pas d'image dans le document. "

    if not resize_OK:
        raisons += "L'image n'est pas redimensionnée de manière à respecter la proportion originale; "
        print(raisons)
    return (points, raisons)


#####################
# Récupérer le document

def listerFichiers(nomDossier, extension=".docx"):
    fichiers = os.listdir(nomDossier)
    listeFichiers = []
    for fichier in fichiers:
        if fichier.endswith(extension):
            listeFichiers.append(fichier)
    return (fichiers, listeFichiers)


def generepar(f):
    # Lisez le fichier PDF

    # TODO : crash on some PDF : ex Test-3 (student name inside, non sharable)

    try :
        with open(f, 'rb') as file:
            # Créez un objet PdfFileReader
            pdf = PyPDF2.PdfReader(file)
            # Obtenez les informations du document
            info = pdf.metadata
            # Vérifiez si le fichier a été généré à partir d'un fichier Word ou LibreOffice
            if info.producer != None:
                if 'Word' in info.producer:
                    return ("Word")
                elif 'LibreOffice' in info.producer:
                    return ("LibreOffice")
                else:
                    return (info.producer)
            else:
                return ("Unknown")
    except Exception as e:
        print('problème avec pdf dans generepar', e)
        # verifier_images_redimensionnees_correctement_PDF(file, pdf, 4)
    finally:
        return ("Unknown")


def verifDocumentWord(filename, word, student, total_pages):
    max = 0
    group = "Unknown"
    to_check_manually = ""
    key="nomFichiers"

    fw = filename[0:-4] + ".docx"
    print("Fichier Word ; ", os.path.abspath(fw))
    try:
        document_pywin32 = word.Documents.Open(os.path.abspath(fw))
    except Exception as e:
        sys.stderr.write("pas de fichier Word ? " + str(e))
        to_check_manually += "vérifier présence fichier docx ! "
        student.to_check.add(key)
        return (max, group, to_check_manually)

    add_word_macro(document_pywin32)



    # check weight
    (student.scores["poids"], student.reasons["poids"]) = verifier_moins_de_3_mo(fw, 3, 2)
    max += 2

    document_pydocx = Document(fw)


    # check Styles
    key = "styles"
    try:
        verifier_styles_word(document_pydocx, student, key)
    except Exception as e:
        to_check_manually += "Styles "
        sys.stderr.write("vérifier style a planté" + str(e))
        student.to_check.add(key)

    # check Quote
    key = "citation"
    try:
        check_quote(document_pydocx, student, key)
    except Exception as e:
        to_check_manually += "citation "
        sys.stderr.write("vérifier citation a planté" + str(e))
        student.to_check.add(key)

    key = "TDM"
    try:
        verifier_TDM_word(document_pywin32, student)
    except Exception as e:
        to_check_manually += "TDM. "
        student.to_check.add(key)
        sys.stderr.write("vérifier TDM a planté" + str(e))

    max += student.max_points[key]

    # check Links and TOC
    key = "lien"
    try:
        check_hyperlinks(py_win32_word_app, student)
    except Exception as e:
        to_check_manually += "Liens. "
        student.to_check.add(key)
        sys.stderr.write("vérifier liens a planté" + str(e))

    max += student.max_points[key]

    # check header and footer
    key="piedDePage"
    try:
        # 2023 - group = check_header_and_footer(py_win32_word_app, student, middle_text_asked="Examen TICE – B1", key="piedDePage")
        group = check_header_and_footer(py_win32_word_app, student, header_to_check, middle_text_asked=middle_footer_to_check, key="piedDePage")
        print_debug(debug, "Groupe : "+group)

        to_check_manually += to_check
    except Exception as e:
        to_check_manually += "en-tête/pieds de page a planté. "
        student.to_check.add(key)
        sys.stderr.write("vérifier en-tête/pieds de page a planté" + str(e))
    student.group = group
    max += student.max_points[key]


    # check return and spaces
    key="espaces"
    try:
        verifier_nombre_enter_et_espaces_word(document_pydocx, student, key=key)
        check_page_returns_word(document_pywin32, student, key)
    except Exception as e:
        to_check_manually += key+" a crashé."
        student.to_check.add(key)
        sys.stderr.write("vérifier espaces/enter a planté" + str(e))

    max += 4

    # check FootNotes
    key = "noteBasPage"
    try:
        check_has_footnotes_word(document_pywin32, stud, key)
    except Exception as e:
        to_check_manually += "note de bas de page"+" a crashé."
        student.to_check.add(key)
        sys.stderr.write("vérifier notes de bas de page a planté" + str(e))

    max += stud.max_points["noteBasPage"]

    key="listes"
    # check lists and sub-lists
    try:
        verifier_listes_word(document_pywin32, student, key)
    except Exception as e:
        to_check_manually += key+" a crashé."
        student.to_check.add(key)
        sys.stderr.write("vérifier listes a planté" + str(e))
    max += student.max_points[key]

    key = "tableau"
    # check table
    try:
        check_tables(py_win32_word_app, student, key)
    except Exception as e:
        to_check_manually += key + " a crashé."
        student.to_check.add(key)
        sys.stderr.write("vérifier "+key+" a planté" + str(e))
    max += student.max_points[key]

    # check images
    key = "images"
    try:
        to_check_manually += "Légende, redimensionnement, texte à gauche. "
        student.to_check.add(key)
        (student.scores["images"], student.reasons["images"]) = verifier_images_redimensionnees_correctement_word(
            document_pywin32, student)
    except Exception as e:
        to_check_manually += key+" a crashé."
        student.to_check.add(key)
        sys.stderr.write("vérifier images a planté" + str(e))
    max += student.max_points[key]

    key="section"
    try:
        check_sections_word(document_pydocx, stud, key)
    except Exception as e:
        to_check_manually += key
        student.to_check.add(key)
        sys.stderr.write("vérifier sections a planté" + str(e))
    max += stud.max_points[key]

    # Ferme le document et Word
    document_pywin32.Close(SaveChanges=False)

    stud.to_check.add("orthographe")
    #stud.to_check.add("")

    return (max, group, to_check_manually)


excel = win32com.client.gencache.EnsureDispatch("Excel.Application")


# TODO attention, si un fichier est ouvert, le programme peut bugger.
# Il "suffit" d'ouvrir le document Word, ça bugge mais fait apparaitre la fenêtre pour "ne pas sauver"...
#           Je ne vois pas encore bien le pourquoi !


def check_pdf_file(f, student):
    to_check_manually = ""
    max = 0
    total_pages = 0
    try :
        with open(f, 'rb') as file:
            # Créez un objet PdfFileReader
            pdf = PyPDF2.PdfReader(file)
            total_pages = verifier_nombre_pages_pdf(pdf, 3, 10, student)
            # check number of pages
            ###(scores_set["pages"], reasons_set["pages"]) = verifier_nombre_pages_PDF(file, pdf 3, 10, 2)
            max += 2
    except Exception as e :
        print('problème avec pdf', e)
            # verifier_images_redimensionnees_correctement_PDF(file, pdf, 4)

    return (max, to_check_manually, total_pages)
# try:
#     wb = excel.Workbooks.Open(excel_file_for_results)
#     wb.Close(SaveChanges=False)
#     print("win32Com Closed the workbook without saving")
# except:
#     print("win32Com Workbook not open")
# try:
#     wb = load_workbook(excel_file_for_results)
#     wb.save(filename=excel_file_for_results, as_template=True)
#     print("pyxl Closed the workbook without saving")
# except:
#     print("pyxl  Workbook not open")
##################################       #########################################
def create_xls_sheets(groups):
    # Créer un nouveau tableur
    workbook = openpyxl.Workbook()
    # Créer une nouvelle feuille par groupe
    worksheets={}
    first_empty_row={}
    # rows={}
    for group in groups:
        worksheet = workbook.create_sheet(group)
        worksheets[group]=worksheet
        first_empty_row[group] = fill_first_lines_excel(worksheet, stud)
       # row[group] = first_empty_row[group]
    #if "Unknown" not in groups:
    #        worksheets.append(workbook.create_sheet("Unknown"))
    workbook.remove(workbook["Sheet"])
    return workbook, worksheets, first_empty_row

def save_in_excel_file(students, groups):
    (workbook,worksheets, first_empty_rows) = create_xls_sheets(groups)
    rows = dict(first_empty_rows)
    for student in students:
        fill_result_line_in_excel(worksheets[student.group], rows[student.group], student)
        rows[student.group] += 1
    for key in groups:
        fill_last_line_in_excel(worksheets[key], rows[key], student, first_empty_rows[key] - 1)

    # Enregistrer le tableur
    execute_ensuring_file_not_open(excel_file_for_results, workbook.save)



if __name__ == "__main__":

    stud = Student()
    groups=["Unknown"]

    py_win32_word_app = win32com.client.Dispatch("Word.Application")

    # on prends la liste des fichiers PDF
    (listefichiers, lf) = listerFichiers(".", ".pdf")
    # print(listefichiers)
    students=[]
    for f in lf:
        stud = Student()

        max = 0
        group = "Unknown"

        # check filename
        # voir par quoi c'est généré : Word ou LibO
        # nom fichier :  2023-01-TIC1—Nom- /2
        verifier_nom_fichiers(f, default_start_of_filename, stud)
        max += 2
        # check 2 formats
        verifier_deux_formats_fichiers(f, listefichiers, 2, stud.scores, stud.reasons)
        max += 2

        # check number of pages
        # (points["pages"], raisons["pages"]) = verifier_nombre_pages_PDF(f, 3, 10, 2)
        # max+=2
        m=0
        tot_pages=0
        to_check=""
        try:
            (m, to_check, tot_pages) = check_pdf_file(f, stud)
        except Exception as e:
            print("problème dans le check_pdf_file : No pdf file ? ")
        max += m
        stud.to_check_manually += to_check
        generateur = generepar(f)

        # check according to OS
        if generateur == "Word":
            (maxPoints, group, to_check) = verifDocumentWord(f, py_win32_word_app, stud, tot_pages)
            stud.to_check_manually += to_check
            max += maxPoints
        elif generateur == "LibreOffice":
            print("Fichier LibreOffice")
        else:
            try:
                (maxPoints, group, to_check) = verifDocumentWord(f, py_win32_word_app, stud, tot_pages)
                stud.to_check_manually += to_check
                max += maxPoints
            except Exception as e:
                sys.stderr.write("ce n'est pas un document Word :-( " + str(e))

        print(stud.firstname, " ", stud.name, " ", group, " : ", str(sum(stud.scores.values())), "sur ", max)
        print("========================================")
        students.append(stud)
        if group not in groups:
            groups.append(group)
        print_debug(debug, str(groups))
        # time.sleep(5)
    py_win32_word_app.Quit()

    # generate xlsx results file
    save_in_excel_file(students, groups)

    print("done")
