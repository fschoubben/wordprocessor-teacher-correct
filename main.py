import sys
import os
# import re

import PyPDF2

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import win32com.client

import openpyxl
# from openpyxl import load_workbook
from openpyxl.comments import Comment
from openpyxl.utils.cell import get_column_letter
from openpyxl.styles import Font, PatternFill

from student import Student

excel_file_for_results = "./2023-01-resultats-automatiques.xlsx"


# TODO  utilisation student systématique
# TODO  give all var and function english names
# TODO  beaucoup plus tard : internationalisation des messages dans le fichier xls

# TODO : vérifier pourquoi Manon Jadot plante chaque fois
# TODO : vérifier si un fichier est déjà ouvert avant de l'ouvrir, en Word ET en excel. Le fermer si oui.
# TODO : ajouter mesure temps par fonction, parce que c'est trope lent (commencer par returns ! )

# TODO : Formats
# todo : orthographe
# todo : citation

# todo pour dans beaucoup plus tard : figer titres et noms

def fill_title_cell_excel(worksheet, value, row, col):
    font = Font(bold=True)
    worksheet.cell(row=row, column=col).value = value
    cell = worksheet[get_column_letter(col) + str(row)]
    cell.font = font


def fill_first_lines_excel(worksheet, student):
    row = 1
    fill_title_cell_excel(worksheet, "Nom", row, 1)
    fill_title_cell_excel(worksheet, "Prénom", row, 2)
    fill_title_cell_excel(worksheet, "Total", row, 3)
    col = 4
    for key in student.scores.keys():
        fill_title_cell_excel(worksheet, key, row=row, col=col)
        col += 1
    fill_title_cell_excel(worksheet, "à vérifier manuellement", row=row, col=col)
    # worksheet.cell(row=row, column=col).value = "à vérifier manuellement"
    row += 1
    worksheet.cell(row=row, column=1).value = "Vérification champs"
    col = 4
    for key, value in student.max_points.items():
        worksheet.cell(row=row, column=col).value = key
        col += 1
    row += 1
    col = 4
    for key, value in student.max_points.items():
        worksheet.cell(row=row, column=col).value = value
        col += 1
    worksheet.cell(row=row, column=3).value = "=sum(" + \
                                              get_column_letter(4) + str(row) + \
                                              ":" + get_column_letter(4 + len(student.scores.items())) + \
                                              str(row) + ")"

    row += 1
    return row


def fill_result_line_in_excel(worksheet, row, student):
    # Pour chaque élément du set, l'ajouter dans une nouvelle cellule
    worksheet.cell(row=row, column=1).value = student.name.capitalize()
    worksheet.cell(row=row, column=2).value = student.firstname.capitalize()
    worksheet.cell(row=row, column=3).value \
        = "=sum(" + get_column_letter(4) + str(row) + ":" + get_column_letter(4 + len(student.scores.items())) + str(
        row) + ")"
    col = 4

    # TODO add conditional formatting : https://openpyxl.readthedocs.io/en/latest/formatting.html
    # Todo add formulas : https://openpyxl.readthedocs.io/en/latest/usage.html?highlight=formula#using-formulae
    for key, value in student.scores.items():
        worksheet.cell(row=row, column=col).value = value
        cell = get_column_letter(col) + str(row)
        blank_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
        worksheet[cell].fill = blank_fill
        # Create the comment with the "Why that score" if any
        why = student.reasons[key]
        if why != "":
            comment = Comment(why, "François Schoubben")
            worksheet[cell].comment = comment
        if key in student.to_check:
            # print("mettre", key, " en jaune")
            yellow_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
            worksheet[cell].fill = yellow_fill
        col += 1

    worksheet.cell(row=row, column=col).value = student.to_check_manually


def fill_last_line_in_excel(worksheet, row, student, number_of_non_student_lines):
    row += 1
    worksheet.cell(row=row, column=1).value = "Moyenne étudiant"
    for col in range(3, len(student.max_points) + 3):
        worksheet.cell(row=row, column=col).value = \
            "=average(" + get_column_letter(col) + str(number_of_non_student_lines + 1) \
            + ":" + get_column_letter(col) + str(row - 1) + ")"

    row += 1
    worksheet.cell(row=row, column=1).value = "Min étudiant"
    for col in range(3, len(student.max_points) + 3):
        worksheet.cell(row=row, column=col).value = \
            "=min(" + get_column_letter(col) + str(number_of_non_student_lines + 1) \
            + ":" + get_column_letter(col) + str(row - 2) + ")"
    row += 1
    worksheet.cell(row=row, column=1).value = "MAX étudiant"
    for col in range(3, len(student.max_points) + 3):
        worksheet.cell(row=row, column=col).value = \
            "=max(" + get_column_letter(col) + str(number_of_non_student_lines + 1) \
            + ":" + get_column_letter(col) + str(row - 3) + ")"

# # gestion des fichiers
# nom fichier :  2023-01-TIC1—Nom-


def verifier_nom_fichiers(mfile, template, student):
    # TODO NextVersion (manually done 2023): vérifier espaces ! (-1)
    raisons = ""
    points = 0
    if mfile.startswith(template):
        points += student.max_points["nomFichiers"]
    # f.replace(format,"")
    mfile = mfile[len(template) + 1:-4]
    mfile = mfile.replace("—", "")
    mfile = mfile.replace(" ", "")
    # if "schoubben" in mfile:  # TODO 2023 : manually done
    #     raisons="mauvais nom de fichier; "
    nomcomplet = mfile.split("-")
    nom = nomcomplet[0]
    prenom = nomcomplet[-1]
    # print("in nom prenom = ", nom, prenom, nomcomplet)
    student.name = nom
    student.firstname = prenom
    student.scores["nomFichiers"] = points
    student.reasons["nomFichiers"] = raisons
    return student.max_points["nomFichiers"]


# 2 formats : word/opendoc + pdf /2 TODO : améliorer pour vérifier types
def verifier_deux_formats_fichiers(filename, liste_fichiers, max_points, scores_set, reasons_set, key="format"):
    # f : filename (string)
    nb_fichiers = 0
    for el in liste_fichiers:
        if filename[0:-4] in el:
            nb_fichiers += 1
    # print("nb fichiers", nbFichiers)
    if nb_fichiers == 2:
        scores_set[key] = max_points
        reasons_set[key] = ""
    else:
        scores_set[key] = 0
        reasons_set[key] = "il n'y a pas les 2 formats de fichier"


# moins de 3Mo (moins de 1Mo) /2
def verifier_moins_de_3_mo(filename, max_size, max_points):
    # f : filename (string)
    try:
        file_info = os.stat(filename)
        if file_info.st_size < max_size * 1000000:
            return max_points, ""
        else:
            return 0, "fichier trop gros"
    except Exception as e:
        sys.stderr.write("erreur de nom de fichier" + str(e))


# Traitement de texte
# min 3, max 10 pages /2


def verifier_nombre_pages_pdf(pdf, min_pages, max_pages, student):
    """ Input : f = fichier pdf existant,
    # min = nombre de page minimum (>=0),
    # max = nombre de pages max (>=0),
    # pts = nombre de points à attribuer
    # output : pts si le nombre de pages de f est compris dans l'intervale [min, max], 0 sinon
    # Obtenez les informations du document"""
    # info = pdf.metadata
    key = "pages"
    doc_pages_count = len(pdf.pages)
    if min_pages <= doc_pages_count <= max_pages:
        student.scores[key] = student.max_points[key]
        student.reasons[key] = ""
    else:
        if len(pdf.pages) > max_pages:
            student.scores[key] = 0
            student.reasons[key] = "Plus de " + str(max_pages) + " pages"
        else:
            student.scores[key] = 0
            student.reasons[key] = "Moins de " + str(min_pages) + " pages"
    return doc_pages_count


# orthographe OK (/0) ?? grammalect ? TODO ?
# utilisation de styles : Titre 1, Titre 2, normal == justifié

def check_sections_word(doc, student, key="section"):
    # CAUTION : these are "bonus points"
    if len(doc.sections) > 1:
        student.scores[key] = 2
        student.reasons[key] = "Bonus : contient plusieurs sections"
    else:
        # CAREFUL: these are "bonus points"
        student.scores[key] = 0
        student.reasons[key] = ""


def lister_styles_word(doc):
    liste_titres = []
    #     for parag in doc.paragraphs:
    #         if ("Heading" in parag.style.name) or ("Titre" in parag.style.name):
    #             #print("oui")
    #             listeTitres.append(parag)
    styles = set()
    for paragraph in doc.paragraphs:
        # Obtenez le nom du style du paragraphe
        style_name = paragraph.style.name
        # Si le nom du style correspond à un style de titre, imprimez le contenu du paragraphe
        styles.add(style_name)
        if style_name in ('Titre 1', 'Titre 2', 'Titre 3', 'Heading 1', 'Heading 2', 'Heading 3'):
            liste_titres.append(paragraph)
    return liste_titres


def verifier_styles_word(doc, student, key="styles"):
    pts = 0
    raison = ""

    # utilisation styles pour titres
    lstyles = lister_styles_word(doc)
    # print("nombre de titres utilisés :", len(lstyles))
    if len(lstyles) > 0:
        pts = student.max_points[key] / 2
    else:
        raison += "Pas de style de titre utilisé. "
    # for el in lstyles:
    #    print(el.style_name, el.text)

    # justifié par le style normal
    normal_style = doc.styles['Normal']
    # print(normal_style.paragraph_format.alignment)
    if normal_style.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
        pts += student.max_points[key] / 2
    else:
        raison += "Le style normal n'est pas justifié. "
        student.to_check_manually += "est-ce que cest bien le style normal qui est utilisé ? si non, est-ce justifié ? "
        student.to_check.add(key)

    student.scores[key] = pts
    student.reasons[key] = raison


def check_page_number_word(doc, total_pages):
    print("Nombre total de pages (Word) : ", total_pages)
    # Récupérer les pieds de page
    footer_text = doc.Sections.Item(1).Footers.Item(1).Range.Text
    if str(total_pages) in footer_text:
        print(f"Le nombre total de pages est présent dans le pied de page : \" ", footer_text, " \"")
    else:
        print("pas de nombre total de pages dans ", footer_text)

    # footers = doc.Sections.Item(1).Footers
    # for footer in footers:
    #     for shape in footer.Shapes:
    #         if shape.Type == 8:
    #             if shape.TextFrame.TextRange.Text.find("{ NUMPAGES }") != -1:
    #                 print("This document contains a field with total number of pages in footer.")
    #                 break
    # else:
    #     print("This document doesn't contain a field with total number of pages in footer.")


# en-tête / pied de page /4
# en haut à droite votre section (NP ou NPS ou NS),
# en bas à gauche votre nom,
# en bas à droite le numéro de page et le nombre de pages,
# en bas et au centre « Examen TICE — B1 » ;
def verifier_entetes_pieds_de_page_word(document, student, total_pages, key="piedDePage"):
    # document : pydocx
    # TODO : vérifier avec champs, parce que Sophie Hodzic a mis un numéro de page, mais sans le texte...
    # TODO : vérifier avec champs, parce que Almina Ajdarpasic a mis le nombre total de page à la main

    # TODO : ne vérifie ni les champs ni les alignements ==> à faire manuellement
    max_points = student.max_points[key]
    raison = ""
    to_check_manually = "pied de page : Nombre total pages ; alignement GMD + en-tête : en haut à droite"
    # group = "unknown"
    points = 0
    # for sec in document.sections:
    #    print(sec)
    section = document.sections[0]
    # nbPages = document.page_count
    # TODO : split in 4 functions : function headers
    header = ""
    i = 0
    for h in section.header.paragraphs:
        header += h.text
        i += 1

    if "NPS" in header:
        points += max_points / 4
        group = "PS"
    elif "NP" in header:
        points += max_points / 4
        group = "NP"
    else:
        raison += "Pas de section écrite correctement en en-tête."
        group = "unknown"
    # print("points après headers", points)

    # TODO : split in 4 functions : function footer_middle
    footer = section.footer.paragraphs[0].text
    if "Examen TICE".lower() in footer.lower() and "B1".lower() in footer.lower():
        points += max_points / 4
        # TODO : vérifier qu'il est bien au milieu
    else:
        raison += "Pas de Examen TICE - B1 écrit au milieu du pied de page."
        # print("milieu OK")
    # print("points après pied de page - milieu", points)

    # TODO : split in 4 functions : function footer_right (and 2 sub-functions : page_number et total_pages)
    # footers =  section.Footers
    if "page" in footer.lower() and ("sur" in footer.lower() or ("de" in footer.lower())):
        points += max_points / 4
    elif "page" in footer.lower():
        points += max_points / 8
        raison += "Pas de nombre total de pages. "
        to_check_manually = "Vérifier nombre de pages."
    else:
        to_check_manually = "Vérifier nombre de pages."
        raison += "Pas de numérotation de pages. "
    #    check_page_number_Word(document_pywin32, total_pages)
    # print("points après pied de page - droite", points)

    # TODO : trouver comment fonctionnent les champs !
    # TODO : vérifier qu'il est bien à droite

    # TODO : split in 4 functions : function footer_left_name
    if (student.name.lower() in footer.lower()) \
            and ((student.firstname.lower() in footer.lower()) or (student.firstname[0].lower() in footer.lower())):
        points += max_points / 4
        # TODO : vérifier qu'il est bien à gauche"
    else:
        raison += "Le nom ne se trouve pas en pied de page."
    # print("points après pied de page - gauche", points)

    # test for aligns
    # foots = footer.split("\t")
    # print(foots)

    student.scores[key] = points
    student.reasons[key] = raison
    student.to_check_manually += to_check_manually
    if student.scores[key] < student.max_points[key]:
        student.to_check.add(key)
    return group


# espaces / sauts de page :
#   max 2 return, max 2 espaces d'affilée, bonus pour style de page/saut de section
def verifier_nombre_enter_et_espaces_word(doc, student, max_return=3, max_spaces=2, key="espaces"):  #
    # doc == pydocx
    points = student.max_points[key]
    # previous_empty = False
    spaces_found = False
    return_reason = ""
    spaces_reasons = ""

    paragraphs = doc.paragraphs
    consecutiveemptyparagraphscount = 0

    for p in paragraphs:
        content = p.text
        content.replace(" ", "")
        content.replace("\t", "")
        if len(content) <= 1:
            consecutiveemptyparagraphscount += 1
            if consecutiveemptyparagraphscount > max_return:
                # more than max_enter return founds
                break
        else:
            consecutiveemptyparagraphscount = 0

    if consecutiveemptyparagraphscount > max_return:
        points -= student.max_points[key] / 2
        return_reason = "Il y a plusieurs retours à la ligne consécutifs dans le document. "

    for paragraph in doc.paragraphs:
        too_many_spaces = " "*max_spaces
        if too_many_spaces in paragraph.text:
            # print("plusieurs espaces")
            if not spaces_found:
                points -= student.max_points[key] / 2
                spaces_found = True
                spaces_reasons = "Il y a plusieurs espaces consécutives dans le document. "
    if points <= 0:
        student.scores[key] = 0
        student.reasons[key] = spaces_reasons + return_reason
        # TODO : créer une classe "fichier word" (objet + méthodes ?)
    else:
        student.scores[key] = points
        if points < student.max_points[key]:
            student.reasons[key] = spaces_reasons + return_reason


def check_page_returns_word(document, student, key):
    # document : win32com
    content = document.Content
    content.Find.ClearFormatting()
    content.Find.Replacement.ClearFormatting()
    content.Find.Text = ""
    content.Find.Format = False
    content.Find.Forward = True
    content.Find.Wrap = 1
    content.Find.Execute(FindText=chr(12))

    if content.Find.Found:
        if student.scores[key] < 3:
            student.scores[key] += 1
            student.reasons[key] += "Il y a des sauts de page"

# listes et sous-listes : /2


def verifier_listes_word(document, student, key):
    listes = False
    sous_listes = False

    # Pour chaque paragraphe du document
    for paragraph in document.Paragraphs:
        # Si le paragraphe est dans une liste
        if paragraph.Range.ListFormat.ListType != 0:
            listes = True
            # Si le paragraphe est dans une sous-liste
            if paragraph.Range.ListFormat.ListLevelNumber > 1:
                sous_listes = True
                break

    if sous_listes:
        student.scores[key] = student.max_points[key]
        student.reasons[key] = ""
    elif listes:
        student.scores[key] = student.max_points[key] / 2
        student.reasons[key] = "Pas de sous-liste"
    else:
        student.scores[key] = 0
        student.reasons[key] = "Pas de liste"


# citation avec guillemets : /2 TODO ??
# note de bas de page avec sources : / 2
def check_has_footnotes_word(doc, student, key="noteBasPage"):
    footnotes = doc.Footnotes
    if footnotes.Count > 0:
        student.scores[key] = student.max_points[key]
        student.reasons[key] = ""
        # print("This document contains footnote(s)")
    else:
        endnotes = doc.Endnotes
        if endnotes.Count > 0:
            student.scores[key] = student.max_points[key]
            student.reasons[key] = "Sous forme de note de fin de document. "
        else:
            student.scores[key] = 0
            student.reasons[key] = "Pas de note de bas de page. "
        # print("This document doesn't contain any footnote")


def check_hyperlinks_and_toc_word(document, student):
    # pywin32
    key_toc = "TDM"
    key_links = "lien"

    # Récupère le texte du document
    text = document.Range().Text
    has_link = False
    has_text_link = False
    # on supprime les TOC du document
    # print("on supprime les ", document.TablesOfContents.Count ," TOC")
    if document.TablesOfContents.Count > 0:
        student.scores[key_toc] = student.max_points[key_toc]
        student.reasons[key_toc] = ""
        toc = document.TablesOfContents(1)
        try:
            toc.Delete()
        except Exception as e:
            sys.stderr.write("supprimer TDM a planté" + str(e))
            student.to_check_manually += "TDM. "
            student.to_check.add(key_toc)

    else:
        student.scores[key_toc] = 0
        student.reasons[key_toc] = "pas de table des matières"

    # for entry in toc.Range.ListFormat.ListString.split("\n"):
    #     print("entry ", entry)
    # print("toc.Range", toc.Range)
    # Pour chaque hyperlien du document
    hyperlinks = ""
    try:
        hyperlinks = document.Range().Hyperlinks
    except Exception as e:
        sys.stderr.write("gestion des hyperliens a planté" + str(e))
        student.to_check_manually += "liens"
        student.to_check.add(key_links)
    if hyperlinks != "":
        for hyperlink in hyperlinks:
            # Récupère le texte de l'hyperlien
            # print('4.1')
            hyperlink_text = ""
            try:
                hyperlink_text = hyperlink.Range.Text
            except Exception as e:
                sys.stderr.write("lecture contenu hyperliens a planté" + str(e) + " " + str(hyperlink) + "\n")
                student.to_check_manually += "Vérifier liens sur mot (crash). "
                student.to_check.add(key_links)
            # TODO : vérifier le nombre ? print("lien trouvé : ", hyperlink_text)
            has_link = True
            if (hyperlink_text in text) and ("http" not in hyperlink_text):
                # print("Le document contient un hyperlien sur le mot '{}'.".format(hyperlink_text))
                has_text_link = True
    if has_text_link:
        student.scores[key_links] = student.max_points[key_links]
        student.reasons[key_links] = ""
    elif has_link:
        student.scores[key_links] = 0
        student.reasons[key_links] = "lien mais pas sur un texte"
    else:
        student.scores[key_links] = 0
        student.reasons[key_links] = "Pas de lien dans le document."


# image : /4 TODO
# * redimenssionnée avec proportions ok
# * légende
# * texte à gauche de l'image
# * espacement de min 0,5 avec le texte
# def verifier_images_redimensionnees_correctement_PDF(file, pdf_reader, maxPoints):
#     pages = pdf_reader.getNumPages()
#
#     for page_number in range(pages):
#         page = pdf_reader.getPage(page_number)
#         for obj in page['/Resources']['/XObject'].values():
#             if obj['/Subtype'] == '/Image':
#                 if '/FlateDecode' in obj['/Filter']:
#                     print("This image is PNG")
#                 elif '/DCTDecode' in obj['/Filter']:
#                     print("This image is JPG")


def verifier_images_redimensionnees_correctement_word(document, student, key="images"):
    # Pour chaque image du document
    raisons = ""
    max_points = student.max_points[key]
    student.to_check_manually += "vérifier si image a gardé proportions + texte à gauche + légende"
    # TODO : vérifier le reste
    resize_ok = True
    points = 0
    has_image = False
    # if len(document.Shapes)>0:
    #     points+=maxPoints/4
    image_count = document.Shapes.Count
    # print("nombre d'images \"Shape\" : ", image_count)
    if image_count > 0:
        has_image = True
        for shape in document.Shapes:
            if shape.Type == 12:  # image
                if shape.HasCaption:
                    print(f"shape : {shape.Name} has a caption")
                else:
                    print(f"shape : {shape.Name} doesn't have a caption")
    image_count = document.InlineShapes.Count
    if image_count > 0:
        has_image = True

    if has_image:
        points += max_points / 4
    else:
        raisons += "Pas d'image dans le document. "

    if not resize_ok:
        raisons += "L'image n'est pas redimensionnée de manière à respecter la proportion originale; "
        print(raisons)
    return points, raisons


#####################
# Récupérer le document

def liste_files(dir_name, extension=".docx"):
    fichiers = os.listdir(dir_name)
    file_list = []
    for fichier in fichiers:
        if fichier.endswith(extension):
            file_list.append(fichier)
    return fichiers, file_list


def generepar(filename):
    # Lisez le fichier PDF

    with open(filename, 'rb') as file:
        # Créez un objet PdfFileReader
        pdf = PyPDF2.PdfReader(file)
        # Obtenez les informations du document
        info = pdf.metadata
        # Vérifiez si le fichier a été généré à partir d'un fichier Word ou LibreOffice
        if not(info.producer is None):
            if 'Word' in info.producer:
                return "Word"
            elif 'LibreOffice' in info.producer:
                return "LibreOffice"
            else:
                return info.producer
        else:
            return "unknown"


def check_word_document(filename, word, student, total_pages):
    max_points = 0
    group = "Unknown"
    to_check_manually = ""
    key = "nomFichiers"

    fw = filename[0:-4] + ".docx"
    print("Fichier Word ; ", fw)
    try:
        document_pywin32 = word.Documents.Open(os.path.abspath(fw))
    except Exception as e:
        sys.stderr.write("pas de fichier Word ? " + str(e))
        to_check_manually += "vérifier présence fichier docx ! "
        student.to_check.add(key)
        return max_points, group, to_check_manually

    # check weight
    (student.scores["poids"], student.reasons["poids"]) = verifier_moins_de_3_mo(fw, 3, 2)
    max_points += 2

    document_pydocx = Document(fw)

    # check Styles
    key = "styles"
    try:
        verifier_styles_word(document_pydocx, student, key)
    except Exception as e:
        to_check_manually += "Styles "
        sys.stderr.write("vérifier style a planté" + str(e))
        student.to_check.add(key)
    max_points += 4

    # check hyperlinks and TOC
    key = "lien"
    try:
        check_hyperlinks_and_toc_word(document_pywin32, student)
    except Exception as e:
        to_check_manually += "Liens et TDM. "
        student.to_check.add(key)
        sys.stderr.write("vérifier liens et TDM a planté" + str(e))

    max_points += 4  # TOC
    max_points += 2  # Links

    # check header and footer
    key = "piedDePage"
    try:
        group = verifier_entetes_pieds_de_page_word(document_pydocx, student, total_pages)
        to_check_manually += to_check
    except Exception as e:
        to_check_manually += "en-tête/pieds de page a planté. "
        student.to_check.add(key)
        sys.stderr.write("vérifier en-tête/pieds de page a planté" + str(e))
    max_points += 4

    # check return and spaces
    key = "espaces"
    try:
        verifier_nombre_enter_et_espaces_word(document_pydocx, student, key=key)
        check_page_returns_word(document_pywin32, student, key)
    except Exception as e:
        to_check_manually += key+" a crashé."
        student.to_check.add(key)
        sys.stderr.write("vérifier espaces/enter a planté" + str(e))

    max_points += 4

    # check FootNotes
    key = "noteBasPage"
    try:
        check_has_footnotes_word(document_pywin32, stud, key)
    except Exception as e:
        to_check_manually += "note de bas de page"+" a crashé."
        student.to_check.add(key)
        sys.stderr.write("vérifier notes de bas de page a planté" + str(e))

    max_points += stud.max_points["noteBasPage"]

    key = "listes"
    # check lists and sub-lists
    try:
        verifier_listes_word(document_pywin32, student, key)
    except Exception as e:
        to_check_manually += key+" a crashé."
        student.to_check.add(key)
        sys.stderr.write("vérifier listes a planté" + str(e))
    max_points += student.max_points[key]

    # check images
    key = "images"
    try:
        to_check_manually += "Légende, redimensionnement, texte à gauche. "
        student.to_check.add(key)
        (student.scores["images"], student.reasons["images"]) = verifier_images_redimensionnees_correctement_word(
            document_pywin32, student)
    except Exception as e:
        to_check_manually += key+" a crashé."
        student.to_check.add(key)
        sys.stderr.write("vérifier images a planté" + str(e))
    max_points += student.max_points[key]

    key = "section"
    try:
        check_sections_word(document_pydocx, stud, key)
    except Exception as e:
        to_check_manually += key
        student.to_check.add(key)
        sys.stderr.write("vérifier sections a planté" + str(e))
    max_points += stud.max_points[key]

    # Ferme le document et Word
    document_pywin32.Close(SaveChanges=False)

    stud.to_check.add("orthographe")
    # stud.to_check.add("")

    return max_points, group, to_check_manually


def check_pdf_file(filename, student):
    to_check_manually = ""
    max_points = 0
    with open(filename, 'rb') as file:
        # Créez un objet PdfFileReader
        pdf = PyPDF2.PdfReader(file)
        total_pages = verifier_nombre_pages_pdf(pdf, 3, 10, student)
        # check number of pages
        max_points += 2

        # verifier_images_redimensionnees_correctement_PDF(file, pdf, 4)
    return max_points, to_check_manually, total_pages


# TODO attention, si un fichier est ouvert, le programme peut bugger.
# Il "suffit" d'ouvrir le document Word, ça bugge mais fait apparaitre la fenêtre pour "ne pas sauver"...
#           Je ne vois pas encore bien le pourquoi !

if __name__ == '__main__':

    excel = win32com.client.gencache.EnsureDispatch("Excel.Application")
    # TODO : check if file is open, then close it before
    # try:
    #     wb = excel.Workbooks.Open(excel_file_for_results)
    #     wb.Close(SaveChanges=False)
    #     print("win32Com Closed the workbook without saving")
    # except:
    #     print("win32Com Workbook not open")
    # try:
    #     wb = load_workbook(excel_file_for_results)
    #     wb.save(filename=excel_file_for_results, as_template=True)
    #     print("pyxl Closed the workbook without saving")
    # except:
    #     print("pyxl  Workbook not open")
    # #################################       #########################################

    # Créer un nouveau tableur
    workbook = openpyxl.Workbook()

    # Créer une nouvelle feuille par groupe
    worksheet_PS = workbook.create_sheet("PS")
    worksheet_NP = workbook.create_sheet("NP")
    worksheet_unknown = workbook.create_sheet("unknown")
    workbook.remove(workbook["Sheet"])

    stud = Student()

    first_empty_row = fill_first_lines_excel(worksheet_NP, stud)
    fill_first_lines_excel(worksheet_PS, stud)
    fill_first_lines_excel(worksheet_unknown, stud)

    row_NP = first_empty_row
    row_PS = first_empty_row
    row_unknown = first_empty_row

    py_win32_word_app = win32com.client.Dispatch("Word.Application")

    # on prend la liste des fichiers PDF
    (listefichiers, lf) = liste_files(".", ".pdf")
    # print(listefichiers)
    for f in lf:
        stud = Student()

        max_score = 0
        grp = "unknown"

        # check filename
        # voir par quoi c'est généré : Word ou LibO
        # nom fichier :  2023-01-TIC1—Nom- /2
        verifier_nom_fichiers(f, "2023-01-TIC1", stud)
        max_score += 2

        # check 2 formats
        verifier_deux_formats_fichiers(f, listefichiers, 2, stud.scores, stud.reasons)
        max_score += 2

        # check number of pages
        # (points["pages"], raisons["pages"]) = verifier_nombre_pages_PDF(f, 3, 10, 2)
        # max+=2
        (m, to_check, tot_pages) = check_pdf_file(f, stud)
        max_score += m
        stud.to_check_manually += to_check
        generateur = generepar(f)

        # check according to OS
        if generateur == "Word":
            (maxPoints, grp, to_check) = check_word_document(f, py_win32_word_app, stud, tot_pages)
            stud.to_check_manually += to_check
            max_score += maxPoints
        elif generateur == "LibreOffice":
            print("Fichier LibreOffice")
        else:
            try:
                (maxPoints, grp, to_check) = check_word_document(f, py_win32_word_app, stud, tot_pages)
                stud.to_check_manually += to_check
                max_score += maxPoints
            except Exception as exc:
                sys.stderr.write("ce n'est pas un document Word :-( " + str(exc))

        print(stud.firstname, " ", stud.name, " ", grp, " : ", str(sum(stud.scores.values())), "sur ", max_score)
        if grp == "NP":
            fill_result_line_in_excel(worksheet_NP, row_NP, stud)
            row_NP += 1
        elif grp == "PS":
            fill_result_line_in_excel(worksheet_PS, row_PS, stud)
            row_PS += 1
        else:
            fill_result_line_in_excel(worksheet_unknown, row_unknown, stud)
            row_unknown += 1
        print("========================================")

    # add stats to spreadsheet file
    fill_last_line_in_excel(worksheet_PS, row_PS, stud, first_empty_row - 1)
    fill_last_line_in_excel(worksheet_NP, row_NP, stud, first_empty_row - 1)
    fill_last_line_in_excel(worksheet_unknown, row_unknown, stud, first_empty_row - 1)
    py_win32_word_app.Quit()

    # save spreadsheet
    workbook.save(excel_file_for_results)

    print("done")
